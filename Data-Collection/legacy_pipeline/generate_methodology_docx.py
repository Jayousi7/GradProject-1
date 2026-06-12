import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Dynamically resolve directories relative to the script location
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
BASE_DIR = os.path.join(PROJECT_ROOT, "legacy_pipeline")

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def build_docx():
    docx_path = os.path.join(BASE_DIR, "ase_data_pipeline_methodology.docx")
    print(f"[*] Starting DOCX generation at: {docx_path}")
    
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ----------------- Document Header -----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("Amman Stock Exchange (ASE) Multi-Era Financial Data Ingestion & ETL Pipeline Methodology")
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy Blue
    title_p.paragraph_format.space_after = Pt(4)
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Data Ingestion, Cleansing, Date Swapped Correction, Ticker Unification, and Local Device Automation (2010–2026)")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x70, 0x80, 0x90) # Slate Grey
    subtitle_p.paragraph_format.space_after = Pt(24)
    
    # Divider line
    p_div = doc.add_paragraph()
    p_div_run = p_div.add_run("―" * 58)
    p_div_run.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)
    p_div.paragraph_format.space_after = Pt(20)

    # ----------------- Section 1 -----------------
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Executive Summary & Architectural Overview")
    h1_run.font.name = 'Calibri Light'
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "This methodology document details the data engineering, ETL (Extract, Transform, Load), quality assurance, "
        "and integration pipelines built to consolidate a continuous, contiguous historical dataset for 8 prominent Jordanian "
        "equities listed on the Amman Stock Exchange (ASE). The pipeline successfully integrates legacy bulk history (2010–2025) "
        "with real-time daily operational records retrieved directly from official exchange servers.\n\n"
        "To ensure complete data integrity, the pipeline addresses significant formatting shifts across eras, fixes date-swapping "
        "export errors, and unifies legacy ticker symbols under active modern symbols. This results in a consolidated 16-column "
        "historical database (combined_historical.csv) covering sixteen years of market transactions."
    )

    # ----------------- Section 2 -----------------
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Target Portfolio & Company Meta-Mappings")
    h2_run.font.name = 'Calibri Light'
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "Through extensive historical mode analysis of the official bulletins, we established standard company codes, names, "
        "and market tiers associated with each equity ticker. These are utilized as a master mapping registry to dynamically "
        "populate daily records starting from 2026, where metadata columns are omitted in raw exchange streams."
    )
    
    # Table of Ticker Mappings
    table = doc.add_table(rows=9, cols=5)
    set_table_borders(table)
    
    headers = ["Ticker", "Company Code", "Arabic Company Name", "Default Market", "Historical / Unification Notes"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        # Font format
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)
                
    ticker_data = [
        ("ARBK", "113023", "البنك العربي", "1", "Arab Bank"),
        ("JOPT", "142041", "مصفاة البترول الأردنية /جوبترول", "1", "Jordan Petroleum Refinery"),
        ("JOEP", "131004", "الكهرباء الاردنية", "1", "Jordan Electric Power"),
        ("JOPH", "141018", "مناجم الفوسفات الاردنية", "1", "Jordan Phosphate Mines"),
        ("UBSI", "111007", "بنك الإتحاد", "1", "Bank al Etihad"),
        ("APOT", "141043", "البوتاس العربية", "1", "Arab Potash"),
        ("RMCC", "141065", "الباطون الجاهز والتوريدات الانشائية", "1", "Ready Mix Concrete"),
        ("ATCO", "141058", "انجاز للتنمية والمشاريع المتعددة", "2", "Unifies LIPO symbol prior to 2016")
    ]
    
    for row_idx, data in enumerate(ticker_data):
        row_cells = table.rows[row_idx + 1].cells
        # Alternate row background
        bg_color = "F4F6F9" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            for p in row_cells[col_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)
                    if col_idx == 0:
                        r.font.bold = True

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(12)
    p_note_run = p_note.add_run(
        "Important Note on ATCO/LIPO Unification: Prior to 2016-01-03, the company انجاز للتنمية والمشاريع المتعددة "
        "(Company Code 141058) traded under the ticker symbol LIPO. On 2016-01-03, the exchange updated its trading symbol to ATCO. "
        "The ETL pipeline programmatically unifies pre-2016 LIPO rows into the ATCO array to present a continuous, sixteen-year timeline."
    )
    p_note_run.font.italic = True
    p_note_run.font.size = Pt(9.5)
    p_note_run.font.color.rgb = RGBColor(0x70, 0x80, 0x90)

    # ----------------- Section 3 -----------------
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Stage 1: Excel-to-CSV Conversion (convert_to_csv.py)")
    h3_run.font.name = 'Calibri Light'
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "The historical bulletins are published in heterogeneous file formats (legacy binary .xls spreadsheets, OpenXML .xlsx spreadsheets, "
        "and tab-delimited text TSV sheets). The convert_to_csv.py script normalizes this raw files storage layer through several core actions:\n\n"
        "•  Format Signature Inspection: The script reads the raw OLE2 byte signatures directly to identify binary structure types (OLE2 vs. Zip-OpenXML) "
        "rather than relying on inaccurate filename extensions.\n"
        "•  Active Tab Selection Logic: Scans raw sheets to discard decorative placeholder tabs and parses the first tab containing valid tabular structures.\n"
        "•  Arabic Text Encoding: TSV text bulletins (used in 2010, 2012, 2015, and 2016) are explicitly decoded using cp1256 (Windows Arabic) parameters "
        "to prevent string corruption of domestic company names.\n"
        "•  Standard Output Format: Writes standard files encoded in UTF-8 with Byte Order Mark (utf-8-sig) to bulletin_staging/raw_csv/<year>/ to ensure identical "
        "rendering in Microsoft Excel, Python, and web browsers."
    )

    # ----------------- Section 4 -----------------
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Stage 2: Ingestion & QA of bulletins (combine_bulletins.py)")
    h4_run.font.name = 'Calibri Light'
    h4_run.font.size = Pt(16)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "The combine_bulletins.py script merges the 16 yearly normalized CSVs into the master combined_bulletins.csv. "
        "It addresses key structural anomalies:"
    )
    
    doc.add_paragraph(
        "A. The 2023–2025 'Swapped Days' Export Corruption\n"
        "In the raw bulletins for the years 2023, 2024, and 2025, the database exporter swapped the day and month fields (rendering 2023-05-03 "
        "instead of 2023-03-05) for any session date where the day of the month was 12 or less. Standard dates after the 12th were formatted correctly.\n"
        "The Solution: The parser, parse_corrupted_dates, identifies dates falling on a day <= 12, splits them, and manually swaps the month and day integers "
        "back to their correct chronological values. All other standard dates are parsed using dayfirst=True parameters.\n\n"
        "B. Column Schema Standardization\n"
        "Standardizes structural column deviations over time. Aligns legacy variations like 2012's trading_date and OPEN_PRICE, and 2021-2022's SEC_CODE, "
        "SEC_NAME1, SYMBOL1, TRADE_QTY, and OPEN_PRICE into standard columns.\n\n"
        "C. Strict Type Alignment & QA Assertions\n"
        "Text columns are stripped of empty spacing. Numeric identifiers are cast to nullable Int64 types to avoid float conversion. Before saving, "
        "the script executes four assertions: Row Count Check (assures zero data loss), Schema Check (assures 16 reference columns), Monotonicity Check "
        "(validates chronological sequence), and Critical Nulls Check (guarantees zero null rows in trade_date, name, code, and symbol)."
    )

    # ----------------- Section 5 -----------------
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Stage 3: Live Daily Local Device Automation (download_and_process.py)")
    h5_run.font.name = 'Calibri Light'
    h5_run.font.size = Pt(16)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "To keep your data continually updated up to today, an automation pipeline runs directly on your local device:\n\n"
        "•  Company Export ID Lookup: The script download_and_process.py parses the company directory HTML to scrape the unique export ID for each ticker "
        "(e.g., JOPH is ID 619).\n"
        "•  Spreadsheet Downloader: Initiates HTTP requests to fetch the live daily historical export (in .xlsx format) directly from ASE servers: "
        "https://ase.com.jo/en/daily-historical-export/{company_id}?_format=xlsx. This query pulls all daily listings up to the absolute current minute.\n"
        "•  Local Command Launcher: The script is launched locally via a Windows batch command pipeline, run_pipeline.bat, which automates file downloads, "
        "cleans raw columns, and immediately triggers the final consolidation mapping script."
    )

    # ----------------- Section 6 -----------------
    h6 = doc.add_paragraph()
    h6_run = h6.add_run("6. Stage 4: Multi-Era Schema Binding (generate_combined.py)")
    h6_run.font.name = 'Calibri Light'
    h6_run.font.size = Pt(16)
    h6_run.font.bold = True
    h6_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "The generate_combined.py script links the pre-2026 bulletins with the 2026 live daily logs. "
        "Starting from 2026, the ASE daily format only records 7 columns because bid/ask quotes and open prices are not captured in the new daily summaries. "
        "We mapped these 7 columns into the standard 16-column schema, filling unrecorded fields with NaN and populating company metadata "
        "(Arabic Names, Numeric Codes, Market Tiers) by looking up their respective modes from the pre-2026 historical bulletins mapping registry."
    )

    # ----------------- Section 7 -----------------
    h7 = doc.add_paragraph()
    h7_run = h7.add_run("7. Unified Output Files and Verification Summary")
    h7_run.font.name = 'Calibri Light'
    h7_run.font.size = Pt(16)
    h7_run.font.bold = True
    h7_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h7.paragraph_format.space_before = Pt(18)
    h7.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "Upon execution, the consolidation pipeline validated and wrote the following outputs, sorted chronologically ascending by Date and Symbol:\n\n"
        "•  Unified Consolidated Master File:\n"
        "   Path: combined_historical.csv\n"
        "   Shape: (29,247 rows, 16 columns)\n"
        "   Date Range: 2010-01-03 to 2026-06-01 (Today)\n\n"
        "•  Individual Company Files (saved inside firms/ folders):\n"
        "   - ARBK_historical.csv\n"
        "   - JOPT_historical.csv\n"
        "   - JOEP_historical.csv\n"
        "   - JOPH_historical.csv\n"
        "   - UBSI_historical.csv\n"
        "   - APOT_historical.csv\n"
        "   - RMCC_historical.csv\n"
        "   - ATCO_historical.csv"
    )
    
    doc.save(docx_path)
    print(f"[+] DOCX successfully created and saved: {docx_path}")

if __name__ == "__main__":
    build_docx()
